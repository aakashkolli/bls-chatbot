#!/usr/bin/env python3
"""Export Qwen v2 20-run results to a Word document (complete outputs, no latencies).

Usage:
  python3 bls_model_eval_pipeline/export_qwen_to_docx.py

This script expects the JSON results file at
`bls_model_eval_pipeline/results/qwen_v2_20_results.json` and will write
`Qwen_v2_20_Run_Output.docx` at the repository root.

It prefers the structured JSON full outputs (single/dual) and falls back to
best-effort extraction if shapes differ.
"""
from pathlib import Path
import json
import sys

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    print("Missing dependency: python-docx. Install with: pip install python-docx")
    raise SystemExit(2)


BASE = Path(__file__).resolve().parent
RESULTS_JSON = BASE / "results" / "qwen_v2_20_results.json"
OUT_DOCX = BASE.parent / "Qwen_v2_20_Run_Output.docx"


def extract_answer(item, key_names):
    """Try a list of key names and shapes to extract a textual answer."""
    for key in key_names:
        if key in item:
            val = item[key]
            if isinstance(val, str):
                return val.strip()
            if isinstance(val, dict):
                # common nested fields
                for sub in ("answer", "text", "output", "final", "response"):
                    if sub in val and isinstance(val[sub], str):
                        return val[sub].strip()
                # stringify fallback
                try:
                    return json.dumps(val, ensure_ascii=False)
                except Exception:
                    return str(val)
            # other types
            try:
                return str(val)
            except Exception:
                continue
    # try common fallback fields
    for f in ("single", "dual", "single_agent", "dual_agent", "refiner", "retriever"):
        if f in item:
            return extract_answer(item, [f])
    return ""


def normalize_entries(data):
    """Return a list of question entries (dicts) from the JSON file."""
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        # Common patterns: {'results': [...]}
        if "results" in data and isinstance(data["results"], list):
            return data["results"]
        # If keys are numeric indices
        numeric_keys = [k for k in data.keys() if k.isdigit()]
        if numeric_keys:
            return [data[k] for k in sorted(numeric_keys, key=int)]
        # If values contain dicts with 'question'
        vals = [v for v in data.values() if isinstance(v, dict) and 'question' in v]
        if vals:
            return vals
        # If top-level is a mapping of single run metadata, try to find list-like fields
        for candidate in ("items", "entries", "rows", "cases"):
            if candidate in data and isinstance(data[candidate], list):
                return data[candidate]
        # fallback: return values that are dicts
        return [v for v in data.values() if isinstance(v, dict)]
    return []


def main():
    if not RESULTS_JSON.exists():
        print(f"Results file not found: {RESULTS_JSON}")
        return 1

    raw = json.loads(RESULTS_JSON.read_text(encoding='utf-8'))
    entries = normalize_entries(raw)
    if not entries:
        print("No entries found in results JSON.")
        return 2

    doc = Document()
    doc.add_heading('Qwen 2.5 v2 — 20-question run (full outputs)', level=1)
    doc.add_paragraph('Source: bls_model_eval_pipeline/results/qwen_v2_20_results.json')

    for i, e in enumerate(entries, start=1):
        # Find question text
        q = e.get('question') or e.get('prompt') or e.get('Question') or e.get('q') or f'Question {i}'
        p = doc.add_paragraph()
        p.add_run(f'{i}. {q}').bold = True

        # Extract single/dual answers using many common key variants
        single = extract_answer(e, ['single', 'single_agent', 'refiner', 'refined_answer', 'answer'])
        dual = extract_answer(e, ['dual', 'dual_agent', 'retriever_refiner', 'final', 'response'])

        # If dual missing but item has 'answers' list with two elements, use that
        if not dual and 'answers' in e and isinstance(e['answers'], list) and len(e['answers']) >= 2:
            dual = e['answers'][1]
            if isinstance(dual, dict):
                dual = extract_answer({'dual': dual}, ['dual'])

        # If single missing but 'answers' present
        if not single and 'answers' in e and isinstance(e['answers'], list) and len(e['answers']) >= 1:
            single = e['answers'][0]
            if isinstance(single, dict):
                single = extract_answer({'single': single}, ['single'])

        # Write Single answer
        doc.add_paragraph('Single-Agent Answer:').bold = True
        doc_para = doc.add_paragraph()
        run = doc_para.add_run(single or '(no single-agent answer)')
        run.font.size = Pt(10)

        # Write Dual answer
        doc.add_paragraph('Dual-Agent Answer:').bold = True
        doc_para = doc.add_paragraph()
        run = doc_para.add_run(dual or '(no dual-agent answer)')
        run.font.size = Pt(10)

        doc.add_paragraph('')

    doc.save(OUT_DOCX)
    print('Wrote', OUT_DOCX)
    return 0


if __name__ == '__main__':
    sys.exit(main())
