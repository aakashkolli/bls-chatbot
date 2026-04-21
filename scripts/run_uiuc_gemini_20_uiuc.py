#!/usr/bin/env python3
"""
Run the 20-question test suite via the UIUC chat API using model "gemini-2.0-flash-lite".
Saves JSON and DOCX outputs to `bls_model_eval_pipeline/results/`.

This script uses the UIUC chat API (UIUC_CHAT_API_KEY) rather than a direct
Gemini API key. It calls the UIUC API endpoint with the model name set to
`gemini-2.0-flash-lite` so the UIUC proxy supplies the appropriate credentials.
"""
from __future__ import annotations

import json
import os
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import requests
from dotenv import load_dotenv

load_dotenv()

UIUC_API_URL = os.getenv("UIUC_API_URL", "https://chat.illinois.edu/api/chat-api/chat")
API_KEY = os.getenv("UIUC_CHAT_API_KEY", "")
COURSE_NAME = os.getenv("COURSE_NAME", "")
MODEL = "gemini-2.0-flash-lite"
TIMEOUT = 120

if not API_KEY:
    raise RuntimeError("Missing UIUC_CHAT_API_KEY in environment (.env)")

# Import system prompts from the app so we match production prompts
from app import SYSTEM_PROMPT_A_V3, SYSTEM_PROMPT_B_V3  # type: ignore

from bls_model_eval_pipeline.questions import get_all_questions_flat  # type: ignore

QUESTION_KEYS_20 = [
    "Prospective Student Questions|General|1",
    "Prospective Student Questions|General|2",
    "Prospective Student Questions|General|4",
    "Prospective Student Questions|Admissions|5",
    "Prospective Student Questions|Admissions|6",
    "Prospective Student Questions|Admissions|8",
    "Prospective Student Questions|Admissions|9",
    "Prospective Student Questions|Academics|11",
    "Prospective Student Questions|Academics|14",
    "Prospective Student Questions|Academics|13",
    "Prospective Student Questions|Cost & Financial Aid|18",
    "Prospective Student Questions|Cost & Financial Aid|20",
    "Prospective Student Questions|Cost & Financial Aid|22",
    "Current Student Questions|Advising & Registration|1",
    "Current Student Questions|Policies|8",
    "More Subjective/Open-Ended Questions||1",
    "More Subjective/Open-Ended Questions||3",
    "More Subjective/Open-Ended Questions||4",
    "More Subjective/Open-Ended Questions||13",
    "More Subjective/Open-Ended Questions||15",
]


def build_question_bank() -> list[dict[str, Any]]:
    by_key = {q["key"]: q for q in get_all_questions_flat()}
    missing = [k for k in QUESTION_KEYS_20 if k not in by_key]
    if missing:
        raise RuntimeError(f"Missing keys in questions.py: {missing}")
    return [by_key[k] for k in QUESTION_KEYS_20]


def extract_text_from_response(obj: Any) -> str:
    if obj is None:
        return ""
    if isinstance(obj, str):
        return obj
    if isinstance(obj, dict):
        # common top-level keys
        if "message" in obj and isinstance(obj["message"], str):
            return obj["message"]
        if "result" in obj and isinstance(obj["result"], str):
            return obj["result"]
        # best-effort traversal
        for k in ("output", "outputs", "choices", "response", "responses", "results", "data", "candidates"):
            if k in obj:
                v = obj[k]
                if isinstance(v, list) and v:
                    first = v[0]
                    if isinstance(first, dict):
                        for tk in ("text", "generated_text", "content", "answer"):
                            if tk in first and isinstance(first[tk], str):
                                return first[tk]
                    elif isinstance(first, str):
                        return first
                elif isinstance(v, str):
                    return v

    # fallback: try to find any string inside
    def find(obj: Any) -> str | None:
        if isinstance(obj, str):
            return obj
        if isinstance(obj, dict):
            for _, v in obj.items():
                found = find(v)
                if found:
                    return found
        if isinstance(obj, list):
            for it in obj:
                found = find(it)
                if found:
                    return found
        return None

    found = find(obj)
    if found:
        return found
    try:
        return json.dumps(obj)
    except Exception:
        return str(obj)


def call_uiuc(system_prompt: str, user_content: str, model: str = MODEL, timeout: int = TIMEOUT) -> dict[str, Any]:
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        "api_key": API_KEY,
        "course_name": COURSE_NAME,
        "stream": False,
        "temperature": 0.1,
        "retrieval_only": False,
    }

    try:
        r = requests.post(UIUC_API_URL, json=payload, headers={"Content-Type": "application/json"}, timeout=timeout)
    except requests.RequestException as exc:
        return {"ok": False, "message": f"API_REQUEST_FAILED: {exc}", "raw": None, "status": 0}

    if r.status_code != 200:
        try:
            j = r.json()
        except Exception:
            return {"ok": False, "message": f"HTTP_{r.status_code}", "raw": r.text, "status": r.status_code}
        return {"ok": False, "message": j.get("error", j), "raw": j, "status": r.status_code}

    try:
        j = r.json()
    except Exception:
        return {"ok": True, "message": r.text or "", "raw": r.text, "status": 200}

    return {"ok": True, "message": j.get("message", j.get("result", j)), "raw": j, "status": 200}


def run() -> None:
    questions = build_question_bank()
    results: dict[str, Any] = {"meta": {"timestamp": datetime.utcnow().isoformat() + "Z", "model": MODEL}, "questions": []}

    for q in questions:
        qobj: dict[str, Any] = {"section": q.get("section"), "subsection": q.get("subsection"), "num": q.get("num"), "text": q.get("text"), "responses": []}

        question_text = q.get("text")
        print(f"[SINGLE] {question_text}")
        single_raw = call_uiuc(SYSTEM_PROMPT_B_V3, question_text)
        single_text = extract_text_from_response(single_raw.get("message") if isinstance(single_raw.get("message"), (str, dict)) else single_raw.get("message"))

        print(f"[RETRIEVER] {question_text}")
        draft_raw = call_uiuc(SYSTEM_PROMPT_A_V3, question_text)
        draft_text = extract_text_from_response(draft_raw.get("message") if isinstance(draft_raw.get("message"), (str, dict)) else draft_raw.get("message"))

        refine_prompt = (
            f"Original User Query: {question_text}\n"
            f"Internal Draft Answer: {draft_text}\n\n"
            "Please refine this draft into a final response following your system constraints."
        )

        print(f"[REFINER] refining draft for: {question_text}")
        dual_raw = call_uiuc(SYSTEM_PROMPT_B_V3, refine_prompt)
        dual_text = extract_text_from_response(dual_raw.get("message") if isinstance(dual_raw.get("message"), (str, dict)) else dual_raw.get("message"))

        qobj["responses"].append({"mode": "single", "text": single_text, "raw": single_raw})
        qobj["responses"].append({"mode": "draft", "text": draft_text, "raw": draft_raw})
        qobj["responses"].append({"mode": "dual", "text": dual_text, "raw": dual_raw})

        results["questions"].append(qobj)

        # Small delay to be polite to the API
        time.sleep(0.8)

    out_dir = Path("bls_model_eval_pipeline") / "results"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_json = out_dir / "gemini_uiuc_runs_20.json"
    with open(out_json, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)

    # Create a simple DOCX report
    try:
        from docx import Document
    except Exception as exc:
        print("python-docx is required to write DOCX. Install with: pip install python-docx")
        raise

    doc = Document()
    doc.add_heading("Gemini (via UIUC) — 20-Question Run", level=1)
    doc.add_paragraph(f"Model: {MODEL}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).isoformat()}")

    for i, q in enumerate(results["questions"], start=1):
        doc.add_heading(f"{i}. {q['text']}", level=2)
        for resp in q["responses"]:
            mode = resp.get("mode")
            text = (resp.get("text") or "(no text)")
            doc.add_paragraph(f"[{mode}]", style="Intense Quote")
            # Keep paragraphs reasonably sized
            for chunk in text.split("\n\n")[:6]:
                doc.add_paragraph(chunk)

    out_docx = out_dir / "gemini_uiuc_runs_20.docx"
    doc.save(out_docx)

    print(f"Wrote JSON: {out_json}")
    print(f"Wrote DOCX: {out_docx}")


if __name__ == "__main__":
    run()
